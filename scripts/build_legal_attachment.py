from pathlib import Path
import shutil
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "legal_submission"
SCREENS = OUT_DIR / "screens"
OUT_KO = OUT_DIR / "잇닿_직업정보제공사업_사전질의_첨부자료.docx"
OUT_EN = OUT_DIR / "ITDAT_JOB_INFO_ATTACHMENT.docx"

TEAL = "0F766E"
LIGHT_TEAL = "F0FDFA"
LIGHT_BLUE = "EFF6FF"
LIGHT_ORANGE = "FFF7ED"
LIGHT_GRAY = "F8FAFC"
BORDER = "CBD5E1"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def set_font(run, size=12, bold=False, color="1F2937"):
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, size=12, bold=False, color="1F2937", space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=2):
    size = 16 if level == 1 else 14
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14 if level == 2 else 0)
    p.paragraph_format.space_after = Pt(7)
    set_font(p.add_run(text), size=size, bold=True, color="0F4C5C")
    if level == 2:
        p_pr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), TEAL)
        borders.append(bottom)
        p_pr.append(borders)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_border(cell, TEAL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(title), size=12, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    set_font(p2.add_run(body), size=12)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.1
        set_font(p.add_run(item), size=12)


def add_steps(doc, items):
    for index, item in enumerate(items, 1):
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(f"{index}단계  "), size=12, bold=True, color=TEAL)
        set_font(p.add_run(item), size=12)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_compact_steps(doc, items):
    table = doc.add_table(rows=0, cols=1)
    table.autofit = False
    for index, item in enumerate(items, 1):
        cell = table.add_row().cells[0]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        set_font(p.add_run(f"{index}. "), size=12, bold=True, color=TEAL)
        set_font(p.add_run(item), size=12)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "E6FFFB")
        set_cell_border(cell)
        if widths:
            cell.width = Cm(widths[i])
        p = cell.paragraphs[0]
        set_font(p.add_run(header), size=12, bold=True, color="134E4A")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i])
            if widths:
                cells[i].width = Cm(widths[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            set_font(p.add_run(value), size=12, bold=(i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_screenshot(doc, filename, caption, explanation):
    path = SCREENS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Cm(5.4))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.line_spacing = 1.15
    set_font(cap.add_run(caption), size=10, color="475569")
    add_callout(doc, "이 화면에서 확인할 수 있는 운영 구조", explanation, LIGHT_GRAY)


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(12)

    doc.core_properties.title = "잇닿 직업정보제공사업 해당 여부 사전질의 첨부자료"
    doc.core_properties.subject = "서비스 운영 구조 설명"
    doc.core_properties.author = "잇닿(Itdat)"

    add_heading(doc, "잇닿 서비스 설명자료", level=1)
    add_text(doc, "직업정보제공사업 해당 여부 사전질의 첨부자료", size=14, bold=True, color="334155")
    add_text(doc, "작성일: 2026년 7월 15일\n사업장 예정 소재지: 경기도 수원시 권선구", size=10, color="64748B")
    add_callout(doc, "잇닿은 무엇인가요?", "병원이 단기근무 정보를 올리면 간호사·간호조무사가 내용을 보고 직접 지원하는 온라인 서비스입니다. 병원과 워커가 서로 선택하며, 근무가 끝난 뒤 임금은 병원이 워커에게 직접 지급합니다.", LIGHT_TEAL)
    add_callout(doc, "자료 작성 목적", "정식 서비스 시작 전에 실제 운영 예정 방식을 설명하고, 관할기관 답변에 따라 필요한 기능을 수정한 뒤 신고 또는 등록 절차를 진행하기 위한 자료입니다.", LIGHT_ORANGE)

    add_heading(doc, "1. 누가 서비스를 사용하나요?")
    add_table(doc, ["이용자", "서비스에서 하는 일"], [
        ["병원", "필요한 근무 날짜, 시간, 장소, 업무, 임금 조건을 직접 작성하고 지원자를 직접 확인합니다."],
        ["워커", "여러 병원의 근무 정보를 직접 비교하고 원하는 공고를 골라 지원합니다."],
        ["잇닿", "병원과 워커가 정보를 입력하고 확인할 수 있는 온라인 공간과 근태·지급관리 도구를 제공합니다."],
    ], [3.2, 12.8])

    add_heading(doc, "2. 서비스 이용 순서")
    add_compact_steps(doc, [
        "병원이 근무 정보를 직접 작성해 게시합니다.",
        "워커가 공고를 검색하고 내용을 비교합니다.",
        "워커가 원하는 공고에 직접 지원합니다.",
        "병원이 지원자 정보를 보고 직접 수락하거나 거절합니다.",
        "병원과 워커가 근로조건을 확인하고 직접 계약합니다.",
        "잇닿에서 출퇴근 시간과 근무시간을 기록합니다.",
        "병원이 워커 본인 명의 계좌로 임금을 직접 지급합니다.",
    ])
    add_callout(doc, "핵심", "잇닿이 특정 워커를 골라 병원에 보내는 방식이 아닙니다. 워커는 공고를 직접 선택하고, 병원은 지원자를 직접 선택합니다.")

    add_heading(doc, "3. 병원에서 보는 실제 화면")
    add_text(doc, "병원은 필요한 자격, 날짜, 시간, 임금과 업무 내용을 직접 입력합니다.")
    add_screenshot(
        doc,
        "admin_shift_new.png",
        "실제 개발 화면 1. 병원 관리자 공고 등록 화면",
        "병원이 근무 자격, 날짜, 시간, 시급, 업무 내용을 직접 정해 입력합니다. 잇닿이 근무조건을 정하거나 특정 워커를 지정하는 항목은 없습니다. 게시된 정보는 워커가 여러 공고를 비교하고 스스로 지원 여부를 판단하기 위한 직업정보로 제공됩니다.",
    )
    add_text(doc, "지원자가 생기면 병원이 지원자의 프로필과 자격정보를 직접 확인한 뒤 수락 또는 거절합니다. 잇닿은 지원자 순위를 정하거나 특정 지원자를 채용하라고 권하지 않습니다.")

    add_heading(doc, "4. 워커가 보는 정보")
    add_table(doc, ["화면", "워커가 확인하거나 하는 일"], [
        ["공고 찾기", "근무 장소, 날짜, 시간, 업무, 자격요건과 병원이 제시한 임금을 비교합니다."],
        ["지원하기", "원하는 공고를 워커가 직접 선택해 지원합니다."],
        ["진행상태", "지원 완료, 병원 확인, 수락·거절 결과를 확인합니다."],
        ["지급상태", "근무 완료, 병원 지급 완료와 본인 계좌 입금 여부를 확인합니다."],
    ], [3.2, 12.8])

    add_heading(doc, "5. 근로계약은 누가 하나요?")
    add_table(doc, ["업무", "담당하는 사람"], [
        ["근무조건과 임금 제시", "병원"],
        ["공고 선택과 지원", "워커"],
        ["지원자 수락 또는 거절", "병원"],
        ["근로조건 확인과 계약", "병원과 워커"],
        ["화면·기록 기능 제공", "잇닿"],
    ], [7.5, 8.5])
    add_text(doc, "잇닿은 근로계약의 당사자, 사용자 또는 대리인이 되지 않습니다. 병원과 워커 사이의 임금이나 근로조건을 대신 협상하지 않습니다.")

    add_heading(doc, "6. 임금은 어떻게 지급하나요?")
    add_steps(doc, [
        "근무시간을 병원과 워커가 확인합니다.",
        "잇닿은 예상 금액 계산과 지급 상태 기록을 돕습니다.",
        "병원이 자기 계좌에서 워커 본인 명의 계좌로 직접 이체합니다.",
        "워커가 입금 여부를 확인합니다.",
    ])
    add_callout(doc, "잇닿은 임금을 다루지 않습니다.", "워커 임금을 받거나 보관하지 않고, 병원을 대신해 지급하지 않으며, 임금에서 플랫폼 수수료를 빼지 않습니다.")
    add_screenshot(
        doc,
        "admin_payroll.png",
        "실제 개발 화면 2. 병원 직접 급여 지급관리 화면",
        "화면에는 병원이 워커에게 임금을 직접 지급한다는 원칙과 잇닿이 임금을 보관하지 않는다는 안내가 표시됩니다. 잇닿은 확인된 근무시간을 바탕으로 예상 지급액과 처리 상태를 기록하는 관리 도구만 제공합니다. 실제 송금 주체는 병원이며, 송금 대상은 워커 본인 명의 계좌입니다.",
    )

    add_heading(doc, "7. 잇닿은 어떻게 수익을 얻나요?")
    add_text(doc, "SaaS란 사업자가 인터넷에서 업무용 소프트웨어를 이용하고 정기적으로 이용료를 내는 방식입니다.")
    add_table(doc, ["과금 항목", "예정 방식"], [
        ["기본 이용료", "병원 또는 지점 단위의 월 정액 이용료"],
        ["추가 사용량", "관리자 수, 활성 워커 수, 알림·API 등 시스템 사용량"],
        ["워커 이용료", "받지 않음"],
        ["채용 성공 수수료", "받지 않음"],
        ["임금 연동 수수료", "받지 않음"],
    ], [5.2, 10.8])
    add_text(doc, "병원이 내는 이용료는 특정 워커의 채용 성공 여부나 임금 액수와 연결되지 않는 소프트웨어 이용료로 운영할 예정입니다.")
    add_screenshot(
        doc,
        "admin_membership.png",
        "실제 개발 화면 3. 워커 임금·채용 성공액과 분리된 SaaS 요금제·청구 화면",
        "병원이 내는 비용은 관리자·근태·지급기록 등 업무용 소프트웨어를 사용하는 월 정액 이용료입니다. 워커의 임금 액수, 지원 횟수, 채용 성사 여부에 따라 소개 수수료가 발생하는 구조가 아닙니다. 워커에게는 가입비나 소개비를 청구하지 않을 예정입니다.",
    )

    add_heading(doc, "8. 잇닿이 하지 않는 일")
    add_bullets(doc, [
        "특정 병원·워커를 골라 연결하거나 지원자를 평가·추천하는 일",
        "임금·근로조건을 중개하거나 계약을 대신하고 당사자가 되는 일",
        "워커 임금을 받거나 보관·대신 지급하는 일",
        "워커에게 가입비·소개비·수수료를 받는 일",
    ])

    add_heading(doc, "9. 관할기관 답변을 받은 뒤의 계획")
    add_compact_steps(doc, [
        "직업정보제공사업이면 필요한 서류를 신고하고 확인증을 받은 뒤 운영합니다.",
        "기능 수정이 필요하면 운영정책과 화면을 먼저 고친 뒤 다시 확인합니다.",
        "유료직업소개사업이면 정식 운영 전에 인력·시설 요건을 갖춰 등록합니다.",
    ])
    add_callout(doc, "제출 목적", "실제 운영 예정 방식을 설명하고, 정식 서비스 전에 필요한 신고 또는 등록 절차를 안내받기 위함입니다.", LIGHT_ORANGE)

    doc.save(OUT_KO)
    shutil.copyfile(OUT_KO, OUT_EN)
    print(OUT_KO)


if __name__ == "__main__":
    build()
